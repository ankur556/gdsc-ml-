import os
import streamlit as st
import PyPDF2
import docx
import chromadb
from chromadb.utils import embedding_functions
from openai import OpenAI
import uuid

# Initialize OpenAI clientimport os
import openai
import docx
import PyPDF2
import speech_recognition as sr
import streamlit as st

# --- Setup ---
st.title("💬 Document-Based Chatbot with Voice & Text")
st.write("This chatbot can help you search and process documents, as well as take voice or text inputs.")

# --- File Upload ---
uploaded_files = st.file_uploader("Upload your documents", accept_multiple_files=True, type=["txt", "pdf", "docx"])

document_chunks = []

# Helper function to read different document formats
def read_text_file(file):
    return file.read().decode("utf-8")

def read_pdf_file(file):
    text = ""
    pdf_reader = PyPDF2.PdfReader(file)
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx_file(file):
    doc = docx.Document(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_document(file, file_extension):
    if file_extension == "txt":
        return read_text_file(file)
    elif file_extension == "pdf":
        return read_pdf_file(file)
    elif file_extension == "docx":
        return read_docx_file(file)
    return ""

# Helper function to split document into chunks
def split_text(text, chunk_size=500):
    sentences = text.replace('\n', ' ').split('. ')
    chunks, current_chunk, current_size = [], [], 0

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if not sentence.endswith('.'):
            sentence += '.'
        sentence_size = len(sentence)

        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size

    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

# --- Process Uploaded Files ---
if uploaded_files:
    document_chunks.clear()
    for file in uploaded_files:
        file_extension = file.name.split('.')[-1].lower()
        content = read_document(file, file_extension)
        chunks = split_text(content)
        document_chunks.extend(chunks)
    st.success(f"Successfully uploaded {len(uploaded_files)} document(s)!")

# --- Simple Search ---
def simple_search(query, n_results=2):
    return [chunk for chunk in document_chunks if query.lower() in chunk.lower()][:n_results]

def get_context_with_sources(results):
    return "\n\n".join(results) if results else "No relevant documents found."

# --- OpenAI Response Generation ---
def generate_response(query, context):
    prompt = f"""Based on the following context, provide a relevant response. If no relevant info is found, say so.
    
    Context:
    {context}
    
    User: {query}
    Assistant:"""

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        max_tokens=500
    )

    return response["choices"][0]["message"]["content"]

# --- Streamlit Chat Interface ---
openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")

if not openai_api_key:
    st.info("Please add your OpenAI API key to proceed.", icon="🔑")
else:
    openai.api_key = openai_api_key  # Set the API key

    # Initialize session state if it doesn't exist
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Display previous messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # User input
    if prompt := st.chat_input("Type your message here..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # Process the input, search for relevant documents, and get the response
        results = simple_search(prompt)
        context = get_context_with_sources(results)
        response = generate_response(prompt, context)

        with st.chat_message("assistant"):
            st.markdown(response)

        # Store assistant's response
        st.session_state.messages.append({"role": "assistant", "content": response})

    # Voice input
    if st.button("Speak"):
        recognizer = sr.Recognizer()
        with sr.Microphone() as source:
            st.write("Listening... Please speak.")
            audio = recognizer.listen(source)
            try:
                voice_input = recognizer.recognize_google(audio)
                st.write(f"You said: {voice_input}")
                st.session_state.messages.append({"role": "user", "content": voice_input})
                with st.chat_message("user"):
                    st.markdown(voice_input)
                results = simple_search(voice_input)
                context = get_context_with_sources(results)
                response = generate_response(voice_input, context)
                with st.chat_message("assistant"):
                    st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
            except sr.UnknownValueError:
                st.write("Sorry, I couldn't understand your speech.")
            except sr.RequestError:
                st.write("Sorry, there was an issue with the speech recognition service.")
client = OpenAI()

# Set up Streamlit UI
st.set_page_config(page_title="RAG Chatbot", layout="wide")
st.title("📚 RAG Chatbot with ChromaDB & OpenAI")

# API Key Input
api_key = st.sidebar.text_input("Enter OpenAI API Key", type="password")
os.environ["OPENAI_API_KEY"] = api_key if api_key else ""

# Upload Files
uploaded_files = st.sidebar.file_uploader("Upload Documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

# Initialize ChromaDB
client = chromadb.PersistentClient(path="Chromadb")
sentence_transformer_ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
collection = client.get_or_create_collection(name="documents_collection", embedding_function=sentence_transformer_ef)

def read_text_file(file):
    return file.read().decode("utf-8")

def read_pdf_file(file):
    text = ""
    pdf_reader = PyPDF2.PdfReader(file)
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx_file(file):
    doc = docx.Document(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_document(file):
    if file.type == "text/plain":
        return read_text_file(file)
    elif file.type == "application/pdf":
        return read_pdf_file(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx_file(file)
    return ""

def process_document(file):
    content = read_document(file)
    return content

if uploaded_files:
    for file in uploaded_files:
        with st.spinner(f"Processing {file.name}..."):
            content = process_document(file)
            chunks = content.split(". ")
            ids = [f"{file.name}_chunk_{i}" for i in range(len(chunks))]
            metadatas = [{"source": file.name, "chunk": i} for i in range(len(chunks))]
            collection.add(documents=chunks, metadatas=metadatas, ids=ids)
            st.success(f"{file.name} added to knowledge base.")

query = st.text_input("Ask a question:")

if st.button("Search") and query:
    results = collection.query(query_texts=[query], n_results=3)
    context = "\n\n".join(results['documents'][0])
    response = client.chat.completions.create(
        model="gpt-4", 
        messages=[{"role": "user", "content": context + "\n\n" + query}],
        temperature=0, max_tokens=500
    ).choices[0].message.content
    st.subheader("Answer:")
    st.write(response)
    st.subheader("Sources:")
    for meta in results['metadatas'][0]:
        st.write(f"{meta['source']} (Chunk {meta['chunk']})")
